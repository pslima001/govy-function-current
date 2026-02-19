from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from typing import List, Tuple

from docx import Document


@dataclass(frozen=True)
class DoctrineRawText:
    paragraphs: List[str]
    text: str


def read_docx_bytes(docx_bytes: bytes) -> DoctrineRawText:
    """
    Extrai texto de DOCX.

    Regras:
      - NÃO usa LLM
      - NÃO interpreta
      - NÃO infere
      - NÃO altera sentido
      - apenas extrai + normaliza whitespace
    """
    if not docx_bytes:
        raise ValueError("docx_bytes vazio ou None")

    doc = Document(BytesIO(docx_bytes))

    paras: List[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        t = " ".join(t.split())
        paras.append(t)

    text = "\n".join(paras).strip()
    return DoctrineRawText(paragraphs=paras, text=text)


def read(source: str | bytes) -> Tuple[str, dict]:
    """Convenience wrapper: aceita path (str) ou bytes.

    Returns:
        (text, {"paragraphs": [...], "source_type": "path"|"bytes"})
    """
    if isinstance(source, bytes):
        raw = read_docx_bytes(source)
        return raw.text, {"paragraphs": raw.paragraphs, "source_type": "bytes"}

    if isinstance(source, str):
        with open(source, "rb") as f:
            data = f.read()
        if not data:
            raise ValueError(f"Arquivo vazio: {source}")
        raw = read_docx_bytes(data)
        return raw.text, {"paragraphs": raw.paragraphs, "source_type": "path"}

    raise ValueError(f"source deve ser str (path) ou bytes, recebeu {type(source)}")
