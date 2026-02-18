"""Tests for govy.doctrine.reader_docx — DOCX text extraction."""

from __future__ import annotations

import io

import pytest
from docx import Document

from govy.doctrine.reader_docx import DoctrineRawText, read_docx_bytes


def _make_docx(paragraphs: list[str]) -> bytes:
    """Build a minimal DOCX in memory from a list of paragraph strings."""
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestReadDocxBytes:
    def test_basic_extraction(self):
        raw = read_docx_bytes(_make_docx(["Hello world", "Second paragraph"]))
        assert isinstance(raw, DoctrineRawText)
        assert raw.paragraphs == ["Hello world", "Second paragraph"]
        assert raw.text == "Hello world\nSecond paragraph"

    def test_whitespace_normalization(self):
        raw = read_docx_bytes(_make_docx(["  multiple   spaces   here  "]))
        assert raw.paragraphs == ["multiple spaces here"]

    def test_empty_paragraphs_skipped(self):
        raw = read_docx_bytes(_make_docx(["First", "", "  ", "Last"]))
        assert raw.paragraphs == ["First", "Last"]

    def test_single_paragraph(self):
        raw = read_docx_bytes(_make_docx(["Only one"]))
        assert len(raw.paragraphs) == 1
        assert raw.text == "Only one"

    def test_empty_docx_returns_empty(self):
        raw = read_docx_bytes(_make_docx([]))
        assert raw.paragraphs == []
        assert raw.text == ""

    def test_invalid_bytes_raises(self):
        with pytest.raises(Exception):
            read_docx_bytes(b"not a docx file")
